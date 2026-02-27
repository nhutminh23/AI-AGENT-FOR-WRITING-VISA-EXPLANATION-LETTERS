from __future__ import annotations


def read_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def read_pdf(path: str) -> str:
    import pdfplumber

    texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            texts.append(text)
    return "\n".join(texts)


def read_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    lines = []

    def push(text: str) -> None:
        t = (text or "").strip()
        if t:
            lines.append(t)

    # Normal paragraphs
    for p in doc.paragraphs:
        push(p.text)

    # Table-based Word forms (common for visa forms)
    def walk_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    push(p.text)
                for nested in cell.tables:
                    walk_table(nested)

    for table in doc.tables:
        walk_table(table)

    # Optional header/footer text
    for section in doc.sections:
        for p in section.header.paragraphs:
            push(p.text)
        for p in section.footer.paragraphs:
            push(p.text)

    # De-duplicate while preserving order (merged cells often repeat)
    unique_lines = list(dict.fromkeys(lines))
    return "\n".join(unique_lines)
