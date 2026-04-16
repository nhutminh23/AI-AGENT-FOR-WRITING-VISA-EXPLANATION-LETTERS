"""
Build Invitation Letter for Australian Visa Application as DOCX.
Template follows EXACT structure from reference PDF.

Dynamic fields (filled from JSON profile):
  - Host info: name, DOB, nationality, passport, address, phone, occupation, income, visa status
  - Guest info: name, DOB, passport, relationship
  - Trip: start_date, end_date, purpose
  - Accompanying persons (if any)
  - Financial commitments

Signature section supports optional auto-insertion of signature image (PNG/JPG).
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def _sanitize(val: Any, fallback: str = "___________") -> str:
    """Return string value or fallback placeholder."""
    if val is None:
        return fallback
    s = str(val).strip()
    return s if s else fallback


def build_invitation_letter_docx(data: Dict[str, Any], signature_image=None) -> io.BytesIO:
    """
    Build the Invitation Letter DOCX from structured data.

    Expected `data` keys:
      host: {full_name, dob, nationality, passport_no, address, phone, occupation, annual_income, visa_status}
      guest: {full_name, dob, passport_no, relationship}
      trip: {start_date, end_date, purpose}
      accompanying: [{full_name, relationship, note}]  (optional)
      financial_commitments: [str]  (optional, defaults provided)
    """
    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Default style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    host = data.get("host", {})
    trip = data.get("trip", {})
    accompanying = data.get("accompanying", [])

    # Normalize: support both "guests" (array) and legacy "guest" (single)
    guests = data.get("guests")
    if not guests and data.get("guest"):
        guests = [data["guest"]]
    if not guests:
        guests = []

    # ============================================================
    # TITLE
    # ============================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INVITATION LETTER FOR AUSTRALIAN VISA APPLICATION")
    run.bold = True
    run.font.size = Pt(14)
    title.paragraph_format.space_after = Pt(18)

    # ============================================================
    # TO:
    # ============================================================
    doc.add_paragraph("To: The Australian Consulate,")

    # ============================================================
    # HOST INFO
    # ============================================================
    host_info = doc.add_paragraph()
    host_info.add_run(f"My name is: {_sanitize(host.get('full_name'))}").bold = False
    host_info.add_run(f"\nDate of Birth: {_sanitize(host.get('dob'))}")
    host_info.add_run(f"\nNationality: {_sanitize(host.get('nationality'), 'Vietnamese')}")
    host_info.add_run(f"\nPassport Number: {_sanitize(host.get('passport_no'))}")

    # ============================================================
    # HOST LIVING/WORKING DETAILS
    # ============================================================
    visa_status = _sanitize(
        host.get("visa_status"),
        "living and working in Australia"
    )
    doc.add_paragraph(
        f"I am currently {visa_status} with the following details:"
    )

    # Bullet list for host details
    details = [
        f"Address: {_sanitize(host.get('address'))}",
        f"Phone number: {_sanitize(host.get('phone'))}",
        f"Occupation: {_sanitize(host.get('occupation'))}",
        f"Annual income: {_sanitize(host.get('annual_income'))}",
    ]
    for d in details:
        p = doc.add_paragraph(d, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    # Visa status paragraph (optional extra line)
    visa_note = host.get("visa_note")
    if visa_note:
        doc.add_paragraph(visa_note)

    # ============================================================
    # GUEST(S) INFO — "I am writing to invite..."
    # ============================================================
    multi = len(guests) > 1

    if multi:
        # Multi-guest: "I am writing to invite the following persons:"
        rel_list = ", ".join(
            f"my {_sanitize(g.get('relationship'), 'relative')}" for g in guests
        )
        doc.add_paragraph(f"I am writing to invite {rel_list}:")
        for g in guests:
            p = doc.add_paragraph()
            p.add_run(f"• {_sanitize(g.get('full_name'))}").bold = False
            p.add_run(f"\n  Date of Birth: {_sanitize(g.get('dob'))}")
            p.add_run(f"\n  Passport Number: {_sanitize(g.get('passport_no'))}")
            rel = _sanitize(g.get("relationship"), "")
            if rel:
                p.add_run(f"\n  Relationship: {rel}")
    else:
        # Single guest (original format)
        guest = guests[0] if guests else {}
        relationship = _sanitize(guest.get("relationship"), "relative")
        invite_para = doc.add_paragraph()
        invite_para.add_run(f"I am writing to invite my {relationship}: ")
        run_name = invite_para.add_run(_sanitize(guest.get("full_name")))
        run_name.bold = False
        invite_para.add_run(f"\nDate of Birth: {_sanitize(guest.get('dob'))}")
        invite_para.add_run(f"\nPassport Number: {_sanitize(guest.get('passport_no'))}")

    # ============================================================
    # TRIP DETAILS
    # ============================================================
    start_date = _sanitize(trip.get("start_date"))
    end_date = _sanitize(trip.get("end_date"))
    purpose = _sanitize(
        trip.get("purpose"),
        "the purpose of family visit and tourism"
    )

    doc.add_paragraph(
        f"to visit me in Australia for a short stay from {start_date} to {end_date} for {purpose}."
    )

    # ============================================================
    # ACCOMPANYING PERSONS (optional)
    # ============================================================
    if accompanying:
        parts = []
        for person in accompanying:
            name = _sanitize(person.get("full_name"))
            rel = _sanitize(person.get("relationship"), "")
            note = person.get("note", "")
            part = f"{name}"
            if rel:
                part = f"{rel}: {name}"
            if note:
                part += f" ({note})"
            parts.append(part)

        if len(parts) == 1:
            acc_text = parts[0]
        else:
            acc_text = ", ".join(parts[:-1]) + f" and {parts[-1]}"

        subj = "they" if multi else f"my {_sanitize(guests[0].get('relationship'), 'relative')}"
        doc.add_paragraph(
            f"During this trip, {subj} will be accompanied by {acc_text}."
        )

    # ============================================================
    # FINANCIAL COMMITMENTS — "I hereby confirm that:"
    # ============================================================
    doc.add_paragraph("I hereby confirm that:")

    if multi:
        # Gender-neutral for group
        them, their, they = "them", "their", "they"
        guests_label = "my guests"
    else:
        g = guests[0] if guests else {}
        gender = _guess_gender(g)
        them = "him" if gender == "M" else "her"
        their = "his" if gender == "M" else "her"
        they = "he" if gender == "M" else "she"
        guests_label = f"my {_sanitize(g.get('relationship'), 'relative')}"

    default_commitments = [
        f"I will cover all travel expenses for {guests_label}, including airfare, accommodation, living expenses, and any other related costs.",
        f"I will provide accommodation for {them} at my residence throughout {their} stay in Australia.",
        f"I will take full financial responsibility and provide support for {them} during the entire visit.",
        f"I will ensure that {they} comply with all Australian laws and visa conditions and leave Australia before {their} visa expires." if multi else f"I will ensure that {they} complies with all Australian laws and visa conditions and leaves Australia before {their} visa expires.",
    ]

    commitments = data.get("financial_commitments") or default_commitments
    for c in commitments:
        p = doc.add_paragraph(c, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    # ============================================================
    # CLOSING STATEMENTS
    # ============================================================
    doc.add_paragraph(
        "This visit is purely for family reunion and short-term tourism purposes, "
        "with no intention to overstay or engage in any unlawful employment."
    )

    visa_word = "visas" if multi else "a visa"
    guests_word = f"{guests_label}" if not multi else "my guests"
    doc.add_paragraph(
        f"I kindly request the Consulate to consider and grant {visa_word} for {guests_word}."
    )

    doc.add_paragraph("Thank you for your time and consideration.")

    # ============================================================
    # SIGNATURE BLOCK (with optional signature image)
    # ============================================================
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(18)
    closing.add_run("Yours sincerely,")

    # Insert signature image if provided, otherwise leave blank space
    if signature_image is not None:
        sig_para = doc.add_paragraph()
        sig_para.paragraph_format.space_before = Pt(6)
        sig_para.paragraph_format.space_after = Pt(0)
        sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            # Read image bytes (supports Flask FileStorage or file-like objects)
            if hasattr(signature_image, 'read'):
                sig_bytes = signature_image.read()
                if hasattr(signature_image, 'seek'):
                    signature_image.seek(0)  # reset for potential re-use
            else:
                sig_bytes = signature_image
            sig_stream = io.BytesIO(sig_bytes)
            run_img = sig_para.add_run()
            run_img.add_picture(sig_stream, width=Inches(2.0))
            logger.info("Signature image inserted into invitation letter")
        except Exception as e:
            logger.warning(f"Failed to insert signature image: {e}")
            # Fallback: blank space
            sig_para.paragraph_format.space_before = Pt(36)
    else:
        # Blank space for manual signature
        sig_space = doc.add_paragraph()
        sig_space.paragraph_format.space_before = Pt(36)
        sig_space.paragraph_format.space_after = Pt(0)

    # Host name under signature
    name_para = doc.add_paragraph()
    run = name_para.add_run(_sanitize(host.get("full_name")).upper())
    run.bold = True
    run.font.size = Pt(13)

    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def _guess_gender(guest: Dict[str, Any]) -> str:
    """Guess gender from sex field or name patterns. Returns 'M' or 'F'."""
    sex = (guest.get("sex") or guest.get("gender") or "").strip().upper()
    if sex in ("M", "MALE", "NAM"):
        return "M"
    if sex in ("F", "FEMALE", "NỮ", "NU"):
        return "F"
    # Guess from Vietnamese name patterns
    name = (guest.get("full_name") or "").upper()
    female_indicators = ["THI ", "NGOC ", "MY ", "LINH", "HUONG", "PHUONG", "HANH", "TRANG"]
    for ind in female_indicators:
        if ind in name:
            return "F"
    return "M"  # default


def render_invitation_letter_text(data: Dict[str, Any]) -> str:
    """
    Render the invitation letter as plain text for UI preview / editing.
    Uses the same structure as build_invitation_letter_docx but outputs string.
    """
    host = data.get("host", {})
    trip = data.get("trip", {})
    accompanying = data.get("accompanying", [])

    # Normalize: support both "guests" (array) and legacy "guest" (single)
    guests = data.get("guests")
    if not guests and data.get("guest"):
        guests = [data["guest"]]
    if not guests:
        guests = []

    multi = len(guests) > 1

    # Pronouns
    if multi:
        them, their, they = "them", "their", "they"
        guests_label = "my guests"
    else:
        g = guests[0] if guests else {}
        gender = _guess_gender(g)
        them = "him" if gender == "M" else "her"
        their = "his" if gender == "M" else "her"
        they = "he" if gender == "M" else "she"
        guests_label = f"my {_sanitize(g.get('relationship'), 'relative')}"

    visa_status = _sanitize(host.get("visa_status"), "living and working in Australia")
    start_date = _sanitize(trip.get("start_date"))
    end_date = _sanitize(trip.get("end_date"))
    purpose = _sanitize(trip.get("purpose"), "the purpose of family visit and tourism")

    lines = []
    lines.append("INVITATION LETTER FOR AUSTRALIAN VISA APPLICATION")
    lines.append("")
    lines.append("To: The Australian Consulate,")
    lines.append("")

    # Host info
    lines.append(f"My name is: {_sanitize(host.get('full_name'))}")
    lines.append(f"Date of Birth: {_sanitize(host.get('dob'))}")
    lines.append(f"Nationality: {_sanitize(host.get('nationality'), 'Vietnamese')}")
    lines.append(f"Passport Number: {_sanitize(host.get('passport_no'))}")
    lines.append("")

    lines.append(f"I am currently {visa_status} with the following details:")
    lines.append(f"  • Address: {_sanitize(host.get('address'))}")
    lines.append(f"  • Phone number: {_sanitize(host.get('phone'))}")
    lines.append(f"  • Occupation: {_sanitize(host.get('occupation'))}")
    lines.append(f"  • Annual income: {_sanitize(host.get('annual_income'))}")
    lines.append("")

    # Visa note
    visa_note = host.get("visa_note")
    if visa_note:
        lines.append(visa_note)
        lines.append("")

    # Guest(s) info
    if multi:
        rel_list = ", ".join(
            f"my {_sanitize(g.get('relationship'), 'relative')}" for g in guests
        )
        lines.append(f"I am writing to invite {rel_list}:")
        lines.append("")
        for g in guests:
            lines.append(f"• {_sanitize(g.get('full_name'))}")
            lines.append(f"  Date of Birth: {_sanitize(g.get('dob'))}")
            lines.append(f"  Passport Number: {_sanitize(g.get('passport_no'))}")
            rel = _sanitize(g.get("relationship"), "")
            if rel:
                lines.append(f"  Relationship: {rel}")
            lines.append("")
    else:
        guest = guests[0] if guests else {}
        relationship = _sanitize(guest.get("relationship"), "relative")
        lines.append(f"I am writing to invite my {relationship}: {_sanitize(guest.get('full_name'))}")
        lines.append(f"Date of Birth: {_sanitize(guest.get('dob'))}")
        lines.append(f"Passport Number: {_sanitize(guest.get('passport_no'))}")
        lines.append("")

    lines.append(f"to visit me in Australia for a short stay from {start_date} to {end_date} for {purpose}.")
    lines.append("")

    # Accompanying
    if accompanying:
        parts = []
        for person in accompanying:
            name = _sanitize(person.get("full_name"))
            rel = _sanitize(person.get("relationship"), "")
            note = person.get("note", "")
            part = f"{name}"
            if rel:
                part = f"{rel}: {name}"
            if note:
                part += f" ({note})"
            parts.append(part)

        if len(parts) == 1:
            acc_text = parts[0]
        else:
            acc_text = ", ".join(parts[:-1]) + f" and {parts[-1]}"

        subj = "they" if multi else guests_label
        lines.append(f"During this trip, {subj} will be accompanied by {acc_text}.")
        lines.append("")

    # Financial commitments
    lines.append("I hereby confirm that:")
    comply_text = f"{they} comply with all Australian laws and visa conditions and leave Australia before {their} visa expires." if multi else f"{they} complies with all Australian laws and visa conditions and leaves Australia before {their} visa expires."
    default_commitments = [
        f"I will cover all travel expenses for {guests_label}, including airfare, accommodation, living expenses, and any other related costs.",
        f"I will provide accommodation for {them} at my residence throughout {their} stay in Australia.",
        f"I will take full financial responsibility and provide support for {them} during the entire visit.",
        f"I will ensure that {comply_text}",
    ]
    commitments = data.get("financial_commitments") or default_commitments
    for c in commitments:
        lines.append(f"  • {c}")
    lines.append("")

    # Closing
    lines.append("This visit is purely for family reunion and short-term tourism purposes, with no intention to overstay or engage in any unlawful employment.")
    lines.append("")
    visa_word = "visas" if multi else "a visa"
    gw = "my guests" if multi else guests_label
    lines.append(f"I kindly request the Consulate to consider and grant {visa_word} for {gw}.")
    lines.append("")
    lines.append("Thank you for your time and consideration.")
    lines.append("")
    lines.append("Yours sincerely,")
    lines.append("")
    lines.append("")
    lines.append(_sanitize(host.get("full_name")).upper())

    return "\n".join(lines)


