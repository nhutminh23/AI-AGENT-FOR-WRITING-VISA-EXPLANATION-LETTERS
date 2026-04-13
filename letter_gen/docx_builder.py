"""
Build DOCX files from generated letter text.
Uses python-docx to create professional A4 documents.
"""
from __future__ import annotations

import io

import logging
import re
import uuid
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def _sanitize_filename(name: str) -> str:
    """Remove special chars from filename."""
    clean = re.sub(r'[^\w\s\-]', '', name)
    clean = re.sub(r'\s+', '_', clean.strip())
    return clean[:60] if clean else "letter"


def build_letter_docx(
    letter_text: str,
    applicant_name: str,
    output_dir: str | Path | None = None,
    filename_prefix: str = "Explanation_Letter",
) -> Path | io.BytesIO:
    """
    Build a DOCX file from letter text.

    Args:
        letter_text: The generated letter text (plain text).
        applicant_name: For filename.
        output_dir: Directory to save the file.
        filename_prefix: Prefix for the filename.

    Returns:
    Returns:
        Path to the saved DOCX file if output_dir is provided, else io.BytesIO.
    """
    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # -- Page setup: A4, 1-inch margins --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # -- Default style --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

    # -- Parse and add paragraphs --
    lines = letter_text.split('\n')
    
    first_title_processed = False

    for i, line in enumerate(lines):
        stripped = line.strip()

        if not stripped:
            # Empty line → add spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        p = doc.add_paragraph()

        # Main Title (The very first non-empty line)
        if not first_title_processed:
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            first_title_processed = True
            continue

        # Detect section headers (e.g., "Purpose of Visit", "Financial Capacity")
        # Heuristic: short line (<60 chars) that doesn't end with period/comma
        is_header = (
            len(stripped) < 80
            and not stripped.endswith(('.', ',', ':', ';'))
            and not stripped.startswith(('•', '-', '–', '—'))
            and i > 0
            and (i == 0 or not lines[i-1].strip())  # preceded by blank line
        )

        # Special cases for known headers
        header_keywords = [
            'Purpose of Visit', 'Travel Arrangements', 'Financial Capacity',
            'Strong Ties', 'Travel History', 'Current Employment',
            'Strong Family Ties', 'Purpose of the Trip', 'Intention to Return',
            'Explanation of Previous', 'SUBJECT:', 'Dear Visa Officer',
            'Key evidence', 'Supporting documents', 'Changed circumstances',
            'Planned itinerary', 'Refusal',
        ]
        is_known_header = any(kw.lower() in stripped.lower() for kw in header_keywords)

        if is_known_header or (is_header and stripped[0].isupper()):
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(12)
        elif stripped.startswith(('•', '-', '–', '—')):
            # Bullet point
            clean = stripped.lstrip('•-–— ').strip()
            p.style = doc.styles['List Bullet']
            p.add_run(clean)
        elif stripped.startswith(('Yours sincerely', 'Sincerely', 'Thank you')):
            p.paragraph_format.space_before = Pt(12)
            p.add_run(stripped)
        else:
            p.add_run(stripped)

    # -- Save --
    if output_dir:
        safe_name = _sanitize_filename(applicant_name)
        session_id = str(uuid.uuid4())[:8]
        filename = f"{filename_prefix}_{safe_name}_{session_id}.docx"
        filepath = output_dir / filename
        doc.save(str(filepath))
        logger.info("DOCX saved: %s (%d bytes)", filepath, filepath.stat().st_size)
        return filepath
    else:
        # Return in-memory bytes buffer
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
