"""
Build Group Tour Participant List DOCX files.
Mirrors the structure of the reference Word template:
 - Group ID header
 - Table: No. | Full Name | Passport No. | Date of Birth | Sex | Date of Expiry

Also provides utility to detect group applications from JSON profiles.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


def detect_group_participants(profiles: Any) -> List[Dict[str, Any]]:
    """
    Detect group participants from JSON profile data.
    
    Accepts either:
      - A list of profiles (each with 'applicant' key)
      - A dict with numbered keys like {"applicant_1": {...}, "applicant_2": {...}}
      - A single profile with 'accompanying_persons' or 'group_members'
    
    Returns:
        List of participant dicts with: full_name, passport_no, dob, sex, passport_expiry
        Returns empty list if < 2 participants detected.
    """
    participants = []

    if isinstance(profiles, list):
        for p in profiles:
            if isinstance(p, dict) and p.get("applicant"):
                participants.append(_extract_participant(p["applicant"]))
    elif isinstance(profiles, dict):
        # Check if it's a single profile with applicant
        if profiles.get("applicant"):
            applicant = profiles["applicant"]
            participants.append(_extract_participant(applicant))
            # Check for accompanying persons / group members
            for key in ("accompanying_persons", "group_members", "co_applicants"):
                extras = profiles.get(key, [])
                if isinstance(extras, list):
                    for person in extras:
                        if isinstance(person, dict):
                            participants.append(_extract_participant(person))
        else:
            # Dict of profiles: {"person_1": {...}, "person_2": {...}}
            for key, val in profiles.items():
                if isinstance(val, dict) and val.get("applicant"):
                    participants.append(_extract_participant(val["applicant"]))

    # Only return if 2+ participants (group)
    if len(participants) < 2:
        return []
    return participants


def _extract_participant(person: Dict[str, Any]) -> Dict[str, Any]:
    """Extract participant fields from an applicant/person dict."""
    return {
        "full_name": person.get("full_name", ""),
        "passport_no": person.get("passport_no", ""),
        "dob": person.get("dob", ""),
        "sex": person.get("sex", person.get("gender", "")),
        "passport_expiry": person.get("passport_expiry", person.get("date_of_expiry", "")),
    }


def build_group_list_json(
    participants: List[Dict[str, Any]],
    group_id: str = "",
) -> Dict[str, Any]:
    """
    Build the Group Tour Participant List as a JSON object.
    
    Returns:
        {
            "group_id": "Q07VZU",
            "participants": [
                {"no": 1, "full_name": "...", "passport_no": "...", "dob": "...", "sex": "...", "passport_expiry": "..."},
                ...
            ]
        }
    """
    return {
        "group_id": group_id,
        "participants": [
            {
                "no": i + 1,
                "full_name": p.get("full_name", ""),
                "passport_no": p.get("passport_no", ""),
                "dob": p.get("dob", ""),
                "sex": p.get("sex", ""),
                "passport_expiry": p.get("passport_expiry", ""),
            }
            for i, p in enumerate(participants)
        ],
    }


def build_group_docx(
    participants: List[Dict[str, Any]],
    group_id: str = "",
) -> io.BytesIO:
    """
    Build a DOCX file containing the Group Tour Participant List table.
    Returns an in-memory BytesIO stream.
    """
    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Default style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GROUP TOUR PARTICIPANT LIST")
    run.bold = True
    run.font.size = Pt(14)
    title.paragraph_format.space_after = Pt(6)

    # Group ID line
    group_line = doc.add_paragraph()
    group_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = group_line.add_run(f"GROUP ID: {group_id or '___________'}")
    run.bold = True
    run.font.size = Pt(12)
    group_line.paragraph_format.space_after = Pt(16)

    # Table headers
    headers = ["No.", "Full Name", "Passport No.", "Date of Birth", "Sex", "Date of Expiry"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for idx, p in enumerate(participants):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = p.get("full_name", "")
        row_cells[2].text = p.get("passport_no", "")
        row_cells[3].text = p.get("dob", "")
        row_cells[4].text = p.get("sex", "")
        row_cells[5].text = p.get("passport_expiry", "")

        for cell in row_cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10)

    # Set column widths
    col_widths = [Cm(1.2), Cm(5), Cm(3), Cm(3.5), Cm(1.8), Cm(3.5)]
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def build_group_header_text(
    participants: List[Dict[str, Any]],
    group_id: str = "",
    group_label: str = "",
) -> str:
    """
    Build the group application header text to prepend to the explanation letter.
    
    Example output:
        Group Application ID: Q07VZU (Ha Family)
        Applicants in this group:
        Mr Tran Trung Anh – Passport No. B4841361
        Mrs Ngo Ngan Ha – Passport No. C3980690
    """
    lines = []
    
    id_line = f"Group Application ID: {group_id or '___________'}"
    if group_label:
        id_line += f" ({group_label})"
    lines.append(id_line)
    
    lines.append("Applicants in this group:")
    for p in participants:
        name = p.get("full_name", "Unknown")
        passport = p.get("passport_no", "")
        if passport:
            lines.append(f"{name} – Passport No. {passport}")
        else:
            lines.append(name)
    
    return "\n".join(lines)
